"""
Production-grade document loading (Chapter 11, Section III).

Real documents fight back: multi-column PDFs, scanned pages, HTML boilerplate,
huge files. A loader you can trust needs two things the one-liner lacks —
fallback strategies for when the primary method fails, and graceful degradation
so one unreadable file does not crash the whole batch.

Heavy/optional libraries (pdfminer, pytesseract, pdf2image, docx, bs4) are
imported lazily inside the methods that use them, so importing this module never
requires the optional extras.
"""
import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


@dataclass
class LoaderConfig:
    clean_whitespace: bool = True
    min_page_chars:   int  = 50      # skip near-empty pages (blank/headers-only)
    max_pages:        int  = 500     # safety limit


class RobustPDFLoader:
    """
    PDF loader with a three-step fallback ladder, each step more powerful and
    more expensive than the last:

        1. pypdf      — fast, handles most PDFs
        2. pdfminer   — better at complex/multi-column layouts
        3. OCR        — last resort for scanned (image-only) PDFs

    Each page becomes its own Document with rich metadata, so retrieved results
    can cite an exact page.
    """

    def __init__(self, config: LoaderConfig | None = None):
        self.config = config or LoaderConfig()

    def load(self, file_path: str) -> list[Document]:
        return list(self._load_with_fallback(Path(file_path)))

    def lazy_load(self, file_path: str) -> Iterator[Document]:
        """Stream pages — preferred for very large documents (flat memory)."""
        yield from self._load_with_fallback(Path(file_path))

    def _load_with_fallback(self, path: Path) -> Iterator[Document]:
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        try:
            yield from self._load_pypdf(path)
            return
        except Exception as exc:
            logger.warning("pypdf failed for %s: %s. Trying pdfminer.", path.name, exc)

        try:
            yield from self._load_pdfminer(path)
            return
        except Exception as exc:
            logger.warning("pdfminer failed for %s: %s. Trying OCR.", path.name, exc)

        try:
            yield from self._load_with_ocr(path)
        except Exception as exc:
            logger.error("All loading strategies failed for %s: %s", path.name, exc)
            # Never crash the pipeline — emit a placeholder and move on.
            yield Document(
                page_content=f"[LOADING FAILED: {path.name}]",
                metadata={"source": str(path), "error": str(exc), "page": 0},
            )

    def _load_pypdf(self, path: Path) -> Iterator[Document]:
        import re
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        total = min(len(reader.pages), self.config.max_pages)
        for i in range(total):
            text = reader.pages[i].extract_text() or ""
            if self.config.clean_whitespace:
                text = re.sub(r"\n{3,}", "\n\n", text)
                text = re.sub(r"[ \t]+", " ", text)
            if len(text) < self.config.min_page_chars:
                continue
            yield Document(
                page_content=text,
                metadata={"source": str(path), "filename": path.name,
                          "page": i + 1, "total_pages": total, "loader": "pypdf"},
            )

    def _load_pdfminer(self, path: Path) -> Iterator[Document]:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer

        for layout in extract_pages(str(path)):
            parts = [el.get_text() for el in layout if isinstance(el, LTTextContainer)]
            text = " ".join(parts).strip()
            if len(text) < self.config.min_page_chars:
                continue
            yield Document(
                page_content=text,
                metadata={"source": str(path), "filename": path.name,
                          "page": getattr(layout, "pageid", None), "loader": "pdfminer"},
            )

    def _load_with_ocr(self, path: Path) -> Iterator[Document]:
        try:
            import pytesseract
            from pdf2image import convert_from_path
        except ImportError as exc:
            raise RuntimeError(
                "OCR requires: pip install pytesseract pdf2image (plus the Tesseract "
                "and poppler system binaries)."
            ) from exc

        images = convert_from_path(str(path), dpi=300)
        for i, image in enumerate(images[: self.config.max_pages], 1):
            text = pytesseract.image_to_string(image, lang="eng")
            if len(text) < self.config.min_page_chars:
                continue
            yield Document(
                page_content=text,
                metadata={"source": str(path), "filename": path.name,
                          "page": i, "loader": "ocr_tesseract"},
            )


class MultiFormatLoader:
    """Routes each file to a format-appropriate loader based on its extension."""

    SUPPORTED = {
        ".pdf": "pdf",
        ".docx": "docx", ".doc": "docx",
        ".html": "html", ".htm": "html",
        ".txt": "text", ".md": "text",
        ".csv": "csv",
    }

    def __init__(self, config: LoaderConfig | None = None):
        self.config = config or LoaderConfig()
        self._pdf = RobustPDFLoader(self.config)

    def load(self, file_path: str) -> list[Document]:
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED:
            raise ValueError(
                f"Unsupported format '{ext}'. Supported: {sorted(self.SUPPORTED)}"
            )
        return getattr(self, f"_load_{self.SUPPORTED[ext]}")(file_path)

    def _load_pdf(self, path: str) -> list[Document]:
        return self._pdf.load(path)

    def _load_docx(self, path: str) -> list[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(path)
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(page_content=text,
                         metadata={"source": path, "filename": Path(path).name,
                                   "loader": "python-docx"})]

    def _load_html(self, path: str) -> list[Document]:
        from langchain_community.document_loaders import BSHTMLLoader
        return BSHTMLLoader(path, open_encoding="utf-8").load()

    def _load_text(self, path: str) -> list[Document]:
        text = Path(path).read_text(encoding="utf-8", errors="replace")
        return [Document(page_content=text,
                         metadata={"source": path, "filename": Path(path).name,
                                   "loader": "text"})]

    def _load_csv(self, path: str) -> list[Document]:
        from langchain_community.document_loaders.csv_loader import CSVLoader
        return CSVLoader(path).load()
